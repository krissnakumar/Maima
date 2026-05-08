
import os
import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

class WordExporter:
    def __init__(self, doc_path="Maima_Captures.docx"):
        self.doc_path = os.path.abspath(doc_path)
        self._ensure_document_exists()

    def _ensure_document_exists(self):
        if not os.path.exists(self.doc_path):
            doc = Document()
            doc.add_heading('Maima Captures', 0)
            doc.save(self.doc_path)

    def add_capture(self, image_path, settings=None):
        settings = settings or {}
        timestamp_toggle = settings.get("timestamp_toggle", True)
        page_layout = settings.get("page_layout", "Portrait")
        image_scaling = settings.get("image_scaling", 100) / 100.0

        try:
            doc = Document(self.doc_path)
            
            # Set Orientation
            section = doc.sections[0]
            if page_layout == "Landscape":
                section.orientation = 1 # WD_ORIENTATION.LANDSCAPE
                # Swap width and height if needed or just set landscape
            
            # Add separator
            doc.add_paragraph("_" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add timestamp
            if timestamp_toggle:
                now = datetime.datetime.now().strftime("%d %b %Y — %H:%M:%S")
                p = doc.add_paragraph()
                run = p.add_run(f"Captured: {now}")
                run.font.size = Pt(10)
                run.italic = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add Image
            # Base width is 6 inches
            base_width = 6.0
            doc.add_picture(image_path, width=Inches(base_width * image_scaling))
            
            # Center the image
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.save(self.doc_path)
            return True
        except Exception as e:
            print(f"Error saving to Word: {e}")
            return False

    def open_document(self):
        try:
            os.startfile(self.doc_path) if hasattr(os, 'startfile') else os.system(f'xdg-open "{self.doc_path}"')
        except Exception as e:
            print(f"Error opening document: {e}")
